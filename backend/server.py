from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, Form, Header, Cookie, Request, Response
from fastapi.responses import FileResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional
import uuid
from datetime import datetime, timezone, timedelta
from passlib.context import CryptContext
from jose import JWTError, jwt
import razorpay
from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContentWithMimeType
import phonenumbers
from phonenumbers import NumberParseException
import json
import random
import PyPDF2
from docx import Document as DocxDocument
from PIL import Image
import io
import base64

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# JWT settings
SECRET_KEY = os.environ.get('JWT_SECRET_KEY', 'your-secret-key-change-in-production')
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_DAYS = 7

# Razorpay client
razorpay_client = razorpay.Client(
    auth=(os.environ.get('RAZORPAY_KEY_ID', ''), os.environ.get('RAZORPAY_KEY_SECRET', ''))
)

# Emergent LLM Key
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY', 'sk-emergent-736B924DfA44eB0D93')

# File upload directory
UPLOAD_DIR = Path("/app/uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# In-memory OTP storage (in production, use Redis)
otp_storage = {}

# Define Models
class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: Optional[str] = None
    name: str
    phone: Optional[str] = None
    password_hash: Optional[str] = None
    auth_provider: str = "email"  # email, google, phone
    profile_picture: Optional[str] = None
    age: Optional[int] = None
    subscription_plan: str = "free"  # free, monthly, yearly
    subscription_status: str = "inactive"  # active, inactive, expired
    subscription_end_date: Optional[datetime] = None
    credits: int = 10  # For metered usage (changed from 0 to 10)
    total_usage: int = 0
    theme: str = "dark"  # dark, light
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class UserSignup(BaseModel):
    email: EmailStr
    password: str
    name: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class PhoneAuth(BaseModel):
    phone: str
    
class PhoneVerify(BaseModel):
    phone: str
    otp: str

class Document(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    session_id: Optional[str] = None  # None means global, otherwise session-specific
    filename: str
    file_type: str
    file_path: str
    file_size: int
    content_preview: Optional[str] = None
    uploaded_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    is_exam_prep: bool = False
    is_global: bool = True  # If True, accessible across sessions

class StudySession(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    type: str  # exam_prep, qa, homework
    name: str
    config: dict = {}  # Stores module-specific config (exam name, date, etc)
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    
class ChatMessage(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    role: str  # user, assistant
    content: str
    sources: List[dict] = []  # For citations
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class StudyMaterial(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    document_id: str
    type: str  # summary, flashcard, mindmap, qa
    content: dict
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Subscription(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    plan_type: str  # monthly, yearly
    amount: int
    status: str
    razorpay_order_id: Optional[str] = None
    razorpay_payment_id: Optional[str] = None
    start_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    end_date: datetime
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class Session(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    session_token: str
    expires_at: datetime
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

# Helper functions
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict):
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + timedelta(days=ACCESS_TOKEN_EXPIRE_DAYS)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(authorization: Optional[str] = Header(None), session_token: Optional[str] = Cookie(None)):
    credentials_exception = HTTPException(
        status_code=401,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    
    token = None
    if session_token:
        # Check session from cookie first
        session = await db.sessions.find_one({"session_token": session_token})
        if session:
            expires_at = session['expires_at']
            if isinstance(expires_at, str):
                expires_at = datetime.fromisoformat(expires_at)
            if expires_at > datetime.now(timezone.utc):
                user = await db.users.find_one({"id": session['user_id']}, {"_id": 0})
                if user:
                    return User(**user)
    
    if authorization and authorization.startswith("Bearer "):
        token = authorization.split(" ")[1]
    
    if not token:
        raise credentials_exception
    
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id: str = payload.get("sub")
        if user_id is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    
    user = await db.users.find_one({"id": user_id}, {"_id": 0})
    if user is None:
        raise credentials_exception
    return User(**user)

async def check_credits(user: User, required_credits: int = 1):
    if user.subscription_status == "active":
        subscription_end = user.subscription_end_date
        if subscription_end and subscription_end > datetime.now(timezone.utc):
            return True
    
    if user.credits < required_credits:
        raise HTTPException(
            status_code=402, 
            detail=f"Insufficient credits. Need {required_credits} credits (₹{required_credits * 5}). Please purchase more credits or subscribe."
        )
    return True

async def deduct_credits(user_id: str, credits: int = 1):
    await db.users.update_one(
        {"id": user_id},
        {"$inc": {"credits": -credits, "total_usage": credits}}
    )

def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract text from PDF, DOCX, TXT, or image files"""
    try:
        if file_type == "application/pdf":
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                return text[:5000]  # Limit preview
        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = DocxDocument(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text[:5000]
        
        elif file_type == "text/plain":
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()[:5000]
        
        return "Text extraction not supported for this file type"
    except Exception as e:
        return f"Error extracting text: {str(e)}"

# Routes
@api_router.get("/")
async def root():
    return {"message": "StudySage API - AI-Powered Study Assistant"}

# Session Management
@api_router.post("/sessions", response_model=Session)
async def create_session(
    request: Request,
    current_user: User = Depends(get_current_user)
):
    data = await request.json()
    session = Session(
        user_id=current_user.id,
        type=data.get('type'),
        name=data.get('name'),
        config=data.get('config', {})
    )
    
    session_dict = session.model_dump()
    session_dict['created_at'] = session_dict['created_at'].isoformat()
    session_dict['updated_at'] = session_dict['updated_at'].isoformat()
    
    await db.sessions_data.insert_one(session_dict)
    return session

@api_router.get("/sessions", response_model=List[Session])
async def get_sessions(
    type: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    query = {"user_id": current_user.id}
    if type:
        query["type"] = type
    
    sessions = await db.sessions_data.find(query, {"_id": 0}).sort("updated_at", -1).to_list(100)
    
    for session in sessions:
        if isinstance(session['created_at'], str):
            session['created_at'] = datetime.fromisoformat(session['created_at'])
        if isinstance(session['updated_at'], str):
            session['updated_at'] = datetime.fromisoformat(session['updated_at'])
    
    return sessions

@api_router.get("/sessions/{session_id}", response_model=Session)
async def get_session(
    session_id: str,
    current_user: User = Depends(get_current_user)
):
    session = await db.sessions_data.find_one({"id": session_id, "user_id": current_user.id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    if isinstance(session['created_at'], str):
        session['created_at'] = datetime.fromisoformat(session['created_at'])
    if isinstance(session['updated_at'], str):
        session['updated_at'] = datetime.fromisoformat(session['updated_at'])
    
    return Session(**session)

@api_router.patch("/sessions/{session_id}")
async def update_session(
    session_id: str,
    request: Request,
    current_user: User = Depends(get_current_user)
):
    data = await request.json()
    update_data = {"updated_at": datetime.now(timezone.utc).isoformat()}
    
    if 'name' in data:
        update_data['name'] = data['name']
    if 'config' in data:
        update_data['config'] = data['config']
    
    result = await db.sessions_data.update_one(
        {"id": session_id, "user_id": current_user.id},
        {"$set": update_data}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Session not found")
    
    return {"message": "Session updated successfully"}

@api_router.delete("/sessions/{session_id}")
async def delete_session(
    session_id: str,
    current_user: User = Depends(get_current_user)
):
    # Delete session
    result = await db.sessions_data.delete_one({"id": session_id, "user_id": current_user.id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Delete associated chat messages
    await db.chat_messages.delete_many({"session_id": session_id})
    
    return {"message": "Session deleted successfully"}

# Chat Messages
@api_router.post("/sessions/{session_id}/messages", response_model=ChatMessage)
async def create_message(
    session_id: str,
    request: Request,
    current_user: User = Depends(get_current_user)
):
    # Verify session belongs to user
    session = await db.sessions_data.find_one({"id": session_id, "user_id": current_user.id})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    data = await request.json()
    message = ChatMessage(
        session_id=session_id,
        role=data.get('role'),
        content=data.get('content'),
        sources=data.get('sources', [])
    )
    
    message_dict = message.model_dump()
    message_dict['created_at'] = message_dict['created_at'].isoformat()
    
    await db.chat_messages.insert_one(message_dict)
    
    # Update session timestamp
    await db.sessions_data.update_one(
        {"id": session_id},
        {"$set": {"updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    return message

@api_router.get("/sessions/{session_id}/messages", response_model=List[ChatMessage])
async def get_messages(
    session_id: str,
    current_user: User = Depends(get_current_user)
):
    # Verify session belongs to user
    session = await db.sessions_data.find_one({"id": session_id, "user_id": current_user.id})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    messages = await db.chat_messages.find({"session_id": session_id}, {"_id": 0}).sort("created_at", 1).to_list(1000)
    
    for msg in messages:
        if isinstance(msg['created_at'], str):
            msg['created_at'] = datetime.fromisoformat(msg['created_at'])
    
    return messages

# Authentication Routes
@api_router.post("/auth/signup")
async def signup(user_data: UserSignup):
    # Check if user exists
    existing_user = await db.users.find_one({"email": user_data.email})
    if existing_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    # Create user
    user = User(
        email=user_data.email,
        name=user_data.name,
        password_hash=get_password_hash(user_data.password),
        auth_provider="email",
        credits=10  # Welcome credits
    )
    
    user_dict = user.model_dump()
    user_dict['created_at'] = user_dict['created_at'].isoformat()
    if user_dict['subscription_end_date']:
        user_dict['subscription_end_date'] = user_dict['subscription_end_date'].isoformat()
    
    await db.users.insert_one(user_dict)
    
    # Create access token
    access_token = create_access_token(data={"sub": user.id})
    
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "user": {"id": user.id, "email": user.email, "name": user.name}
    }

@api_router.post("/auth/login")
async def login(user_data: UserLogin):
    user = await db.users.find_one({"email": user_data.email}, {"_id": 0})
    if not user or not user.get('password_hash'):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    if not verify_password(user_data.password, user['password_hash']):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    access_token = create_access_token(data={"sub": user['id']})
    
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "user": {"id": user['id'], "email": user['email'], "name": user['name']}
    }

@api_router.post("/auth/phone/send-otp")
async def send_otp(phone_data: PhoneAuth):
    try:
        parsed = phonenumbers.parse(phone_data.phone, "IN")
        if not phonenumbers.is_valid_number(parsed):
            raise HTTPException(status_code=400, detail="Invalid phone number")
        
        formatted_phone = phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.E164)
        
        # Generate OTP
        otp = str(random.randint(100000, 999999))
        otp_storage[formatted_phone] = {
            "otp": otp,
            "expires_at": datetime.now(timezone.utc) + timedelta(minutes=10)
        }
        
        # In production, send OTP via SMS service
        # For now, return OTP in response (remove in production)
        return {"message": "OTP sent successfully", "phone": formatted_phone, "otp": otp}
    
    except NumberParseException:
        raise HTTPException(status_code=400, detail="Invalid phone number format")

@api_router.post("/auth/phone/verify-otp")
async def verify_otp(verify_data: PhoneVerify):
    try:
        parsed = phonenumbers.parse(verify_data.phone, "IN")
        formatted_phone = phonenumbers.format_number(parsed, phonenumbers.PhoneNumberFormat.E164)
        
        stored_otp = otp_storage.get(formatted_phone)
        if not stored_otp:
            raise HTTPException(status_code=400, detail="OTP not found or expired")
        
        if stored_otp['expires_at'] < datetime.now(timezone.utc):
            del otp_storage[formatted_phone]
            raise HTTPException(status_code=400, detail="OTP expired")
        
        if stored_otp['otp'] != verify_data.otp:
            raise HTTPException(status_code=400, detail="Invalid OTP")
        
        # Remove used OTP
        del otp_storage[formatted_phone]
        
        # Check if user exists
        user = await db.users.find_one({"phone": formatted_phone}, {"_id": 0})
        
        if not user:
            # Create new user
            new_user = User(
                phone=formatted_phone,
                name=f"User {formatted_phone[-4:]}",
                auth_provider="phone",
                credits=10
            )
            user_dict = new_user.model_dump()
            user_dict['created_at'] = user_dict['created_at'].isoformat()
            if user_dict['subscription_end_date']:
                user_dict['subscription_end_date'] = user_dict['subscription_end_date'].isoformat()
            
            await db.users.insert_one(user_dict)
            user = user_dict
        
        access_token = create_access_token(data={"sub": user['id']})
        
        return {
            "access_token": access_token,
            "token_type": "bearer",
            "user": {"id": user['id'], "phone": user['phone'], "name": user['name']}
        }
    
    except NumberParseException:
        raise HTTPException(status_code=400, detail="Invalid phone number format")

@api_router.post("/auth/google/callback")
async def google_callback(request: Request, response: Response):
    data = await request.json()
    session_id = data.get('session_id')
    
    if not session_id:
        raise HTTPException(status_code=400, detail="Session ID required")
    
    # Get session data from Emergent Auth
    import aiohttp
    async with aiohttp.ClientSession() as session:
        async with session.get(
            "https://demobackend.emergentagent.com/auth/v1/env/oauth/session-data",
            headers={"X-Session-ID": session_id}
        ) as resp:
            if resp.status != 200:
                raise HTTPException(status_code=400, detail="Invalid session")
            
            session_data = await resp.json()
    
    # Check if user exists
    user = await db.users.find_one({"email": session_data['email']}, {"_id": 0})
    
    if not user:
        # Create new user
        new_user = User(
            email=session_data['email'],
            name=session_data['name'],
            profile_picture=session_data.get('picture'),
            auth_provider="google",
            credits=10
        )
        user_dict = new_user.model_dump()
        user_dict['created_at'] = user_dict['created_at'].isoformat()
        if user_dict['subscription_end_date']:
            user_dict['subscription_end_date'] = user_dict['subscription_end_date'].isoformat()
        
        await db.users.insert_one(user_dict)
        user = user_dict
    
    # Create session
    session_token = str(uuid.uuid4())
    session_obj = Session(
        user_id=user['id'],
        session_token=session_token,
        expires_at=datetime.now(timezone.utc) + timedelta(days=7)
    )
    
    session_dict = session_obj.model_dump()
    session_dict['expires_at'] = session_dict['expires_at'].isoformat()
    session_dict['created_at'] = session_dict['created_at'].isoformat()
    await db.sessions.insert_one(session_dict)
    
    # Set cookie
    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        max_age=7 * 24 * 60 * 60,
        path="/"
    )
    
    access_token = create_access_token(data={"sub": user['id']})
    
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "user": {"id": user['id'], "email": user['email'], "name": user['name']}
    }

@api_router.get("/auth/me")
async def get_me(current_user: User = Depends(get_current_user)):
    return current_user

@api_router.patch("/auth/me")
async def update_me(
    request: Request,
    current_user: User = Depends(get_current_user)
):
    data = await request.json()
    update_data = {}
    
    if 'age' in data:
        update_data['age'] = data['age']
    if 'theme' in data:
        update_data['theme'] = data['theme']
    if 'name' in data:
        update_data['name'] = data['name']
    
    if update_data:
        await db.users.update_one(
            {"id": current_user.id},
            {"$set": update_data}
        )
    
    return {"message": "Profile updated successfully"}

@api_router.post("/auth/logout")
async def logout(session_token: Optional[str] = Cookie(None), response: Response = None):
    if session_token:
        await db.sessions.delete_one({"session_token": session_token})
    
    if response:
        response.delete_cookie("session_token")
    
    return {"message": "Logged out successfully"}

# Document Management
@api_router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    session_id: Optional[str] = Form(None),
    is_global: bool = Form(True),
    is_exam_prep: bool = Form(False),
    current_user: User = Depends(get_current_user)
):
    # Check file type
    allowed_types = [
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/plain",
        "image/jpeg",
        "image/png"
    ]
    
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="File type not supported")
    
    # If session_id provided, verify it belongs to user
    if session_id:
        session = await db.sessions_data.find_one({"id": session_id, "user_id": current_user.id})
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")
    
    # Save file
    file_id = str(uuid.uuid4())
    file_extension = file.filename.split('.')[-1]
    file_path = UPLOAD_DIR / f"{file_id}.{file_extension}"
    
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)
    
    # Extract text preview
    content_preview = extract_text_from_file(str(file_path), file.content_type)
    
    # Create document record
    document = Document(
        user_id=current_user.id,
        session_id=session_id,
        filename=file.filename,
        file_type=file.content_type,
        file_path=str(file_path),
        file_size=len(content),
        content_preview=content_preview,
        is_exam_prep=is_exam_prep,
        is_global=is_global
    )
    
    doc_dict = document.model_dump()
    doc_dict['uploaded_at'] = doc_dict['uploaded_at'].isoformat()
    
    await db.documents.insert_one(doc_dict)
    
    return document

@api_router.get("/documents", response_model=List[Document])
async def get_documents(
    is_exam_prep: Optional[bool] = None,
    current_user: User = Depends(get_current_user)
):
    query = {"user_id": current_user.id}
    if is_exam_prep is not None:
        query["is_exam_prep"] = is_exam_prep
    
    documents = await db.documents.find(query, {"_id": 0}).to_list(100)
    
    for doc in documents:
        if isinstance(doc['uploaded_at'], str):
            doc['uploaded_at'] = datetime.fromisoformat(doc['uploaded_at'])
    
    return documents

@api_router.delete("/documents/{document_id}")
async def delete_document(
    document_id: str,
    current_user: User = Depends(get_current_user)
):
    document = await db.documents.find_one({"id": document_id, "user_id": current_user.id})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete file
    file_path = Path(document['file_path'])
    if file_path.exists():
        file_path.unlink()
    
    # Delete document record
    await db.documents.delete_one({"id": document_id})
    
    # Delete related study materials
    await db.study_materials.delete_many({"document_id": document_id})
    
    return {"message": "Document deleted successfully"}

# AI Features
@api_router.post("/ai/summarize/{document_id}")
async def summarize_document(
    document_id: str,
    current_user: User = Depends(get_current_user)
):
    await check_credits(current_user, 2)
    
    document = await db.documents.find_one({"id": document_id, "user_id": current_user.id})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Check if summary already exists
    existing_summary = await db.study_materials.find_one({
        "document_id": document_id,
        "user_id": current_user.id,
        "type": "summary"
    })
    
    if existing_summary:
        return existing_summary
    
    # Create AI summary using Gemini (supports files)
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"summary_{document_id}",
        system_message="You are an expert study assistant. Create comprehensive yet concise summaries of documents."
    ).with_model("gemini", "gemini-2.0-flash")
    
    file_content = FileContentWithMimeType(
        file_path=document['file_path'],
        mime_type=document['file_type']
    )
    
    user_message = UserMessage(
        text=f"Create a comprehensive summary of this document '{document['filename']}'. Include key points, main ideas, and important details. Format it in a clear, structured way.",
        file_contents=[file_content]
    )
    
    try:
        response = await chat.send_message(user_message)
        
        # Save summary
        study_material = StudyMaterial(
            user_id=current_user.id,
            document_id=document_id,
            type="summary",
            content={"summary": response, "document_name": document['filename']}
        )
        
        material_dict = study_material.model_dump()
        material_dict['created_at'] = material_dict['created_at'].isoformat()
        await db.study_materials.insert_one(material_dict)
        
        await deduct_credits(current_user.id, 2)
        
        return study_material
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating summary: {str(e)}")

@api_router.post("/ai/flashcards/{document_id}")
async def create_flashcards(
    document_id: str,
    current_user: User = Depends(get_current_user)
):
    await check_credits(current_user, 2)
    
    document = await db.documents.find_one({"id": document_id, "user_id": current_user.id})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Create flashcards using AI
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"flashcards_{document_id}",
        system_message="You are an expert study assistant. Create effective flashcards from documents."
    ).with_model("gemini", "gemini-2.0-flash")
    
    file_content = FileContentWithMimeType(
        file_path=document['file_path'],
        mime_type=document['file_type']
    )
    
    user_message = UserMessage(
        text="Create 10-15 flashcards from this document. Return them in JSON format as an array of objects with 'question' and 'answer' fields. Example: [{\"question\": \"What is...\", \"answer\": \"...\"}, ...]",
        file_contents=[file_content]
    )
    
    try:
        response = await chat.send_message(user_message)
        
        # Parse flashcards
        flashcards = []
        try:
            # Try to extract JSON from response
            import re
            json_match = re.search(r'\[.*\]', response, re.DOTALL)
            if json_match:
                flashcards = json.loads(json_match.group())
        except:
            # Fallback: parse manually
            flashcards = [{"question": "Question 1", "answer": response}]
        
        # Save flashcards
        study_material = StudyMaterial(
            user_id=current_user.id,
            document_id=document_id,
            type="flashcard",
            content={"flashcards": flashcards, "document_name": document['filename']}
        )
        
        material_dict = study_material.model_dump()
        material_dict['created_at'] = material_dict['created_at'].isoformat()
        await db.study_materials.insert_one(material_dict)
        
        await deduct_credits(current_user.id, 2)
        
        return study_material
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating flashcards: {str(e)}")

@api_router.post("/ai/qa")
async def ask_question(
    request: Request,
    current_user: User = Depends(get_current_user)
):
    await check_credits(current_user, 1)
    
    data = await request.json()
    question = data.get('question')
    document_id = data.get('document_id')
    
    if not question:
        raise HTTPException(status_code=400, detail="Question is required")
    
    # Get document if provided
    context_message = question
    file_content = None
    
    if document_id:
        document = await db.documents.find_one({"id": document_id, "user_id": current_user.id})
        if document:
            file_content = FileContentWithMimeType(
                file_path=document['file_path'],
                mime_type=document['file_type']
            )
    
    # Answer question using AI
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"qa_{current_user.id}",
        system_message="You are an expert homework assistant. Provide detailed, educational answers to questions. Explain concepts clearly."
    ).with_model("gemini", "gemini-2.0-flash")
    
    user_message = UserMessage(
        text=context_message,
        file_contents=[file_content] if file_content else []
    )
    
    try:
        response = await chat.send_message(user_message)
        
        await deduct_credits(current_user.id, 1)
        
        return {"answer": response, "question": question}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error answering question: {str(e)}")

@api_router.post("/ai/mindmap/{document_id}")
async def create_mindmap(
    document_id: str,
    current_user: User = Depends(get_current_user)
):
    await check_credits(current_user, 2)
    
    document = await db.documents.find_one({"id": document_id, "user_id": current_user.id})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Create mindmap using AI
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"mindmap_{document_id}",
        system_message="You are an expert study assistant. Create hierarchical mindmaps from documents."
    ).with_model("gemini", "gemini-2.0-flash")
    
    file_content = FileContentWithMimeType(
        file_path=document['file_path'],
        mime_type=document['file_type']
    )
    
    user_message = UserMessage(
        text="Create a mindmap from this document. Return it in JSON format with a hierarchical structure: {\"title\": \"Main Topic\", \"children\": [{\"title\": \"Subtopic 1\", \"children\": [...]}, ...]}",
        file_contents=[file_content]
    )
    
    try:
        response = await chat.send_message(user_message)
        
        # Parse mindmap
        mindmap = {}
        try:
            import re
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                mindmap = json.loads(json_match.group())
        except:
            mindmap = {"title": document['filename'], "children": []}
        
        # Save mindmap
        study_material = StudyMaterial(
            user_id=current_user.id,
            document_id=document_id,
            type="mindmap",
            content={"mindmap": mindmap, "document_name": document['filename']}
        )
        
        material_dict = study_material.model_dump()
        material_dict['created_at'] = material_dict['created_at'].isoformat()
        await db.study_materials.insert_one(material_dict)
        
        await deduct_credits(current_user.id, 2)
        
        return study_material
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating mindmap: {str(e)}")

@api_router.get("/study-materials")
async def get_study_materials(
    document_id: Optional[str] = None,
    type: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    query = {"user_id": current_user.id}
    if document_id:
        query["document_id"] = document_id
    if type:
        query["type"] = type
    
    materials = await db.study_materials.find(query, {"_id": 0}).to_list(100)
    
    for material in materials:
        if isinstance(material['created_at'], str):
            material['created_at'] = datetime.fromisoformat(material['created_at'])
    
    return materials

# Subscription & Payment
@api_router.post("/subscription/create-order")
async def create_subscription_order(
    request: Request,
    current_user: User = Depends(get_current_user)
):
    data = await request.json()
    plan_type = data.get('plan_type')  # monthly or yearly
    
    if plan_type not in ['monthly', 'yearly']:
        raise HTTPException(status_code=400, detail="Invalid plan type")
    
    amount = 39900 if plan_type == 'monthly' else 399900  # in paise
    
    try:
        order = razorpay_client.order.create({
            "amount": amount,
            "currency": "INR",
            "payment_capture": 1
        })
        
        # Create subscription record
        subscription = Subscription(
            user_id=current_user.id,
            plan_type=plan_type,
            amount=amount,
            status="pending",
            razorpay_order_id=order['id'],
            end_date=datetime.now(timezone.utc) + timedelta(days=30 if plan_type == 'monthly' else 365)
        )
        
        sub_dict = subscription.model_dump()
        sub_dict['start_date'] = sub_dict['start_date'].isoformat()
        sub_dict['end_date'] = sub_dict['end_date'].isoformat()
        sub_dict['created_at'] = sub_dict['created_at'].isoformat()
        await db.subscriptions.insert_one(sub_dict)
        
        return {"order_id": order['id'], "amount": amount, "currency": "INR"}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating order: {str(e)}")

@api_router.post("/subscription/verify-payment")
async def verify_payment(
    request: Request,
    current_user: User = Depends(get_current_user)
):
    data = await request.json()
    razorpay_order_id = data.get('razorpay_order_id')
    razorpay_payment_id = data.get('razorpay_payment_id')
    razorpay_signature = data.get('razorpay_signature')
    
    try:
        # Verify signature
        razorpay_client.utility.verify_payment_signature({
            'razorpay_order_id': razorpay_order_id,
            'razorpay_payment_id': razorpay_payment_id,
            'razorpay_signature': razorpay_signature
        })
        
        # Update subscription
        subscription = await db.subscriptions.find_one({
            "razorpay_order_id": razorpay_order_id,
            "user_id": current_user.id
        })
        
        if not subscription:
            raise HTTPException(status_code=404, detail="Subscription not found")
        
        await db.subscriptions.update_one(
            {"id": subscription['id']},
            {"$set": {"status": "active", "razorpay_payment_id": razorpay_payment_id}}
        )
        
        # Update user
        end_date = datetime.fromisoformat(subscription['end_date'])
        await db.users.update_one(
            {"id": current_user.id},
            {"$set": {
                "subscription_plan": subscription['plan_type'],
                "subscription_status": "active",
                "subscription_end_date": end_date.isoformat()
            }}
        )
        
        return {"message": "Payment verified successfully", "status": "active"}
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Payment verification failed: {str(e)}")

@api_router.post("/credits/purchase")
async def purchase_credits(
    request: Request,
    current_user: User = Depends(get_current_user)
):
    data = await request.json()
    credits_amount = data.get('credits', 50)  # Default 50 credits
    
    # ₹1 per credit
    amount = credits_amount * 100  # in paise
    
    try:
        order = razorpay_client.order.create({
            "amount": amount,
            "currency": "INR",
            "payment_capture": 1,
            "notes": {"credits": credits_amount, "user_id": current_user.id}
        })
        
        return {"order_id": order['id'], "amount": amount, "credits": credits_amount}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating order: {str(e)}")

@api_router.post("/credits/verify-payment")
async def verify_credits_payment(
    request: Request,
    current_user: User = Depends(get_current_user)
):
    data = await request.json()
    razorpay_order_id = data.get('razorpay_order_id')
    razorpay_payment_id = data.get('razorpay_payment_id')
    razorpay_signature = data.get('razorpay_signature')
    credits_amount = data.get('credits', 50)
    
    try:
        # Verify signature
        razorpay_client.utility.verify_payment_signature({
            'razorpay_order_id': razorpay_order_id,
            'razorpay_payment_id': razorpay_payment_id,
            'razorpay_signature': razorpay_signature
        })
        
        # Add credits to user
        await db.users.update_one(
            {"id": current_user.id},
            {"$inc": {"credits": credits_amount}}
        )
        
        return {"message": "Credits added successfully", "credits": credits_amount}
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Payment verification failed: {str(e)}")

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()